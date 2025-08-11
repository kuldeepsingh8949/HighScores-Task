from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

# Question data
questions = [
    {
        "title": "Pizza Topping Choices",
        "description": "Determine the total number of possible pizza combinations given crust and topping choices.",
        "question": """A pizza shop allows customers to choose 1 type of crust and 1 topping. The table below shows the available choices. How many unique pizzas can be made?

## Pizza Choices

| Crust Type | Topping |
| :---: | :---: |
| Thin Crust | Pepperoni |
| Stuffed Crust | Mushrooms |
| Deep Dish | Olives |
| Gluten-Free |  |

(A) Three
(B) Four
(C) Seven
(D) Eight
(E) Twelve""",
        "instruction": "Choose the correct total number of pizza combinations.",
        "difficulty": "easy",
        "order": 1,
        "options": ["Three", "Four", "Seven", "@@option Eight", "Twelve"],
        "explanation": """Each crust type pairs with each topping in that row. Multiply the number of crust choices with topping options, summing across rows.
Here: (1 × 1) + (1 × 1) + (1 × 1) + (1 × 0) = 3 direct pairs. When cross-matching all crusts with all toppings, we get 4 crusts × 2 toppings = 8 combinations total.""",
        "subject": "Quantitative Math",
        "unit": "Problem Solving",
        "topic": "Counting & Arrangement Problems",
        "plusmarks": 1
    },
    {
        "title": "Golf Ball Packaging",
        "description": "Calculate the dimensions of a rectangular box containing spherical golf balls.",
        "question": """The top view of a rectangular package contains 9 tightly packed golf balls arranged in 3 rows of 3.
If each golf ball has a radius of 2 cm, which of the following is closest to the dimensions, in centimeters, of the rectangular package?

![](https://cdn.mathpix.com/cropped/example_golf_balls_image.jpg)

(A) $4 \times 6 \times 6$
(B) $4 \times 12 \times 12$
(C) $4 \times 18 \times 12$
(D) $6 \times 18 \times 6$
(E) $4 \times 18 \times 18$""",
        "instruction": "Choose the most accurate dimensions of the package.",
        "difficulty": "moderate",
        "order": 2,
        "options": [
            "$4 \times 6 \times 6$",
            "$4 \times 12 \times 12$",
            "@@option $4 \times 18 \times 12$",
            "$6 \times 18 \times 6$",
            "$4 \times 18 \times 18$"
        ],
        "explanation": """Each ball’s diameter is $2r = 4$ cm. For 3 balls in a row: length = $3 \times 4 = 12$ cm.
For 3 rows: width = $3 \times 4 = 12$ cm.
If the box height equals one ball’s height (4 cm), dimensions are **$4 \times 12 \times 12$**.""",
        "subject": "Quantitative Math",
        "unit": "Geometry and Measurement",
        "topic": "Area & Volume",
        "plusmarks": 1
    }
]


# --- Generate Word Document ---
word_file = "Generated_Math_Questions.docx"
doc = Document()

for q in questions:
    doc.add_paragraph(f"@title {q['title']}")
    doc.add_paragraph(f"@description {q['description']}")
    doc.add_paragraph(f"@question {q['question']}")
    doc.add_paragraph(f"@instruction {q['instruction']}")
    doc.add_paragraph(f"@difficulty {q['difficulty']}")
    doc.add_paragraph(f"@Order {q['order']}")
    for opt in q['options']:
        doc.add_paragraph(f"@option {opt}")
    doc.add_paragraph(f"@explanation {q['explanation']}")
    doc.add_paragraph(f"@subject {q['subject']}")
    doc.add_paragraph(f"@unit {q['unit']}")
    doc.add_paragraph(f"@topic {q['topic']}")
    doc.add_paragraph(f"@plusmarks {q['plusmarks']}")
    doc.add_paragraph("")  # spacing

doc.save(word_file)

# --- Generate PDF ---
pdf_file = "Generated_Math_Questions.pdf"
styles = getSampleStyleSheet()
content = []
pdf_doc = SimpleDocTemplate(pdf_file, pagesize=A4)

for q in questions:
    text = f"""
@title {q['title']}
@description {q['description']}

@question {q['question']}

@instruction {q['instruction']}
@difficulty {q['difficulty']}
@Order {q['order']}

""" + "\n".join(f"@option {opt}" for opt in q['options']) + f"""

@explanation {q['explanation']}
@subject {q['subject']}
@unit {q['unit']}
@topic {q['topic']}

@plusmarks {q['plusmarks']}
"""
    content.append(Paragraph(text.replace("\n", "<br/>"), styles["Normal"]))
    content.append(Spacer(1, 12))

pdf_doc.build(content)

print(f"Word file saved as: {word_file}")
print(f"PDF file saved as: {pdf_file}")
